"""DOCX (Word document) payload generator.

This module implements DOCX-based hiding techniques targeting systems
that process Word documents, including RAG pipelines, document converters,
and text extraction tools.

Techniques:
    DOCX_HIDDEN_TEXT: Payload in text with hidden font attribute.
    DOCX_TINY_TEXT: Payload in 0.5pt font (below human visual threshold).
    DOCX_WHITE_TEXT: White text on white background.
    DOCX_COMMENT: Payload in Word comment/annotation.
    DOCX_METADATA: Payload in document core properties.
    DOCX_HEADER_FOOTER: Payload in document header or footer.

Targets:
    RAG document pipelines, Word-to-text converters, document summarizers,
    enterprise document processing systems, email attachment scanners.

Usage:
    >>> from countersignal.ipi.generators.docx import create_docx
    >>> from countersignal.ipi.models import Technique
    >>> campaign = create_docx(
    ...     Path("./output.docx"),
    ...     Technique.DOCX_HIDDEN_TEXT,
    ...     "http://localhost:8080"
    ... )
"""

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from docx.document import Document as DocumentType

from countersignal.core.models import Campaign
from countersignal.ipi.models import Format, PayloadStyle, PayloadType, Technique

from . import create_campaign_ids, generate_payload

# DOCX techniques
DOCX_TECHNIQUES = [
    Technique.DOCX_HIDDEN_TEXT,
    Technique.DOCX_TINY_TEXT,
    Technique.DOCX_WHITE_TEXT,
    Technique.DOCX_COMMENT,
    Technique.DOCX_METADATA,
    Technique.DOCX_HEADER_FOOTER,
]
"""All DOCX techniques for document processing pipelines."""


# =============================================================================
# Decoy Content
# =============================================================================


def _create_decoy_document() -> "DocumentType":
    """Create a plausible Word document with decoy content.

    Returns:
        Document object with realistic business content.
    """
    doc = Document()

    # Title
    doc.add_heading("Q4 Financial Summary Report", level=0)

    # Introduction
    doc.add_paragraph(
        "This report summarizes the financial performance for the fourth quarter "
        "of the fiscal year. All figures have been reviewed and approved by the "
        "finance department."
    )

    # Section 1
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Revenue for Q4 exceeded expectations by 12%, driven primarily by strong "
        "performance in the enterprise solutions division. Operating margins "
        "improved to 23.5%, up from 21.2% in the previous quarter."
    )

    # Section 2
    doc.add_heading("Key Metrics", level=1)
    doc.add_paragraph(
        "Total Revenue: $47.3 million\n"
        "Operating Income: $11.1 million\n"
        "Net Profit Margin: 18.7%\n"
        "Customer Acquisition Cost: $1,247"
    )

    # Section 3
    doc.add_heading("Outlook", level=1)
    doc.add_paragraph(
        "Management expects continued growth in Q1 of the new fiscal year, "
        "with projected revenue increases of 8-12%. Strategic investments in "
        "cloud infrastructure are expected to yield returns in the second half."
    )

    return doc


# =============================================================================
# Technique Implementations
# =============================================================================


def _inject_hidden_text(doc: "DocumentType", payload: str) -> None:
    """Inject payload as hidden text using font.hidden attribute.

    Hidden text is invisible in normal view but often extracted by
    text parsers that don't respect the hidden attribute.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    para = doc.add_paragraph()
    run = para.add_run(payload)
    run.font.hidden = True


def _inject_tiny_text(doc: "DocumentType", payload: str) -> None:
    """Inject payload in 0.5pt font size.

    Text at 0.5pt is below human visual threshold but still
    extractable by text parsers.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    para = doc.add_paragraph()
    run = para.add_run(payload)
    run.font.size = Pt(0.5)


def _inject_white_text(doc: "DocumentType", payload: str) -> None:
    """Inject payload as white text on white background.

    White text on default white background is invisible to readers
    but extracted as normal text.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    para = doc.add_paragraph()
    run = para.add_run(payload)
    run.font.color.rgb = RGBColor(255, 255, 255)


def _inject_comment(doc: "DocumentType", payload: str) -> None:
    """Inject payload as a Word comment/annotation.

    Comments are often extracted by document processing tools
    even though they're not part of the main content.

    Note: This uses direct XML manipulation as python-docx doesn't
    have built-in comment support.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    # Add a paragraph that will have the comment attached
    para = doc.add_paragraph("Please review the attached figures.")

    # Get the paragraph XML element
    p = para._p

    # Create comment ID
    comment_id = "1"

    # Build the comment XML namespace map
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Create the comment range start marker
    comment_start = parse_xml(f'<w:commentRangeStart xmlns:w="{w_ns}" w:id="{comment_id}"/>')
    p.insert(0, comment_start)

    # Create the comment range end marker
    comment_end = parse_xml(f'<w:commentRangeEnd xmlns:w="{w_ns}" w:id="{comment_id}"/>')
    p.append(comment_end)

    # Create comment reference run
    comment_ref_run = parse_xml(
        f'<w:r xmlns:w="{w_ns}"><w:commentReference w:id="{comment_id}"/></w:r>'
    )
    p.append(comment_ref_run)

    # Create comments part XML content
    comments_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:comments xmlns:w="{w_ns}">'
        f'<w:comment w:id="{comment_id}" w:author="Reviewer" w:date="2024-01-15T10:30:00Z">'
        f"<w:p><w:r><w:t>{payload}</w:t></w:r></w:p>"
        f"</w:comment>"
        f"</w:comments>"
    )

    # Add comments part using OPC package API
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    # Get the document part
    document_part = doc.part

    # Create comments part with proper PackURI
    comments_partname = PackURI("/word/comments.xml")
    comments_content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    )

    # Create the part
    comments_part = Part(
        comments_partname,
        comments_content_type,
        comments_xml.encode("utf-8"),
        document_part.package,
    )

    # Relate the comments part to the document part
    document_part.relate_to(comments_part, RT.COMMENTS)


def _inject_metadata(doc: "DocumentType", payload: str) -> None:
    """Inject payload into document core properties.

    Core properties (author, subject, keywords, etc.) are often
    extracted for indexing and cataloging.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    core_props = doc.core_properties
    core_props.author = payload
    core_props.subject = payload
    core_props.keywords = payload
    core_props.comments = payload


def _inject_header_footer(doc: "DocumentType", payload: str) -> None:
    """Inject payload into document header and footer.

    Headers and footers are extracted by most document parsers
    but may not be prominently displayed to users.

    Args:
        doc: Document to modify.
        payload: Payload string to inject.
    """
    # Access default section
    section = doc.sections[0]

    # Add to header
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_run = header_para.add_run(payload)
    # Make it subtle - small and light gray
    header_run.font.size = Pt(6)
    header_run.font.color.rgb = RGBColor(200, 200, 200)

    # Add to footer
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_run = footer_para.add_run(payload)
    # Make it subtle - small and light gray
    footer_run.font.size = Pt(6)
    footer_run.font.color.rgb = RGBColor(200, 200, 200)


# =============================================================================
# Main DOCX Creation
# =============================================================================


def create_docx(
    output_path: Path,
    technique: Technique,
    callback_url: str,
    payload_style: PayloadStyle = PayloadStyle.OBVIOUS,
    payload_type: PayloadType = PayloadType.CALLBACK,
    seed: int | None = None,
    sequence: int = 0,
) -> Campaign:
    """Generate a DOCX file with hidden prompt injection payload.

    Creates a plausible Word document and injects the payload using
    the specified technique.

    Args:
        output_path: Where to save the DOCX file.
        technique: Hiding technique (see DOCX_TECHNIQUES).
        callback_url: Base URL for callbacks.
        payload_style: Style of payload content (obvious vs subtle).
        payload_type: Objective of the payload.

        seed: Optional seed for deterministic UUID/token generation.
        sequence: Sequence number for batch deterministic generation.

    Returns:
        Campaign object with UUID and metadata.

    Raises:
        ValueError: If technique is not a DOCX technique.

    Example:
        >>> from countersignal.ipi.generators.docx import create_docx
        >>> from countersignal.ipi.models import Technique
        >>> campaign = create_docx(
        ...     Path("./report.docx"),
        ...     Technique.DOCX_HIDDEN_TEXT,
        ...     "http://localhost:8080"
        ... )
    """
    if technique not in DOCX_TECHNIQUES:
        raise ValueError(f"Unsupported DOCX technique: {technique.value}")

    canary_uuid, token = create_campaign_ids(seed, sequence)
    payload = generate_payload(callback_url, canary_uuid, payload_style, payload_type, token=token)

    # Create document with decoy content
    doc = _create_decoy_document()

    # Inject payload using selected technique
    if technique == Technique.DOCX_HIDDEN_TEXT:
        _inject_hidden_text(doc, payload)
    elif technique == Technique.DOCX_TINY_TEXT:
        _inject_tiny_text(doc, payload)
    elif technique == Technique.DOCX_WHITE_TEXT:
        _inject_white_text(doc, payload)
    elif technique == Technique.DOCX_COMMENT:
        _inject_comment(doc, payload)
    elif technique == Technique.DOCX_METADATA:
        _inject_metadata(doc, payload)
    elif technique == Technique.DOCX_HEADER_FOOTER:
        _inject_header_footer(doc, payload)

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return Campaign(
        uuid=canary_uuid,
        token=token,
        filename=output_path.name,
        format=Format.DOCX,
        technique=technique,
        payload_style=payload_style,
        payload_type=payload_type,
        callback_url=callback_url,
    )


# =============================================================================
# Batch Generation
# =============================================================================


def create_all_docx_variants(
    output_dir: Path,
    callback_url: str,
    base_name: str = "report",
    payload_style: PayloadStyle = PayloadStyle.OBVIOUS,
    payload_type: PayloadType = PayloadType.CALLBACK,
    techniques: list[Technique] | None = None,
    seed: int | None = None,
) -> list[Campaign]:
    """Generate DOCX files using multiple techniques.

    Args:
        output_dir: Directory to save files.
        callback_url: Base URL for callbacks.
        base_name: Base filename (technique suffix will be added).
        payload_style: Style of payload content.
        payload_type: Objective of the payload.
        techniques: List of techniques to use (default: all DOCX techniques).

        seed: Optional seed for deterministic UUID/token generation.

    Returns:
        List of Campaign objects.

    Example:
        >>> from countersignal.ipi.generators.docx import create_all_docx_variants
        >>> campaigns = create_all_docx_variants(
        ...     Path("./output"),
        ...     "http://localhost:8080"
        ... )
        >>> len(campaigns)
        6
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    campaigns = []

    if techniques is None:
        techniques = DOCX_TECHNIQUES

    for i, technique in enumerate(techniques):
        filename = f"{base_name}_{technique.value}.docx"
        output_path = output_dir / filename
        campaign = create_docx(
            output_path,
            technique,
            callback_url,
            payload_style,
            payload_type,
            seed=seed,
            sequence=i,
        )
        campaigns.append(campaign)

    return campaigns
